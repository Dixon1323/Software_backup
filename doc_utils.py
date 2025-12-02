import os
import shutil
import uuid
import json
import zipfile
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

from logger import log
from config import TEMPLATE_ORIG, OUTPUT_DIR, LOCAL_DIR
from data_utils import find_shift_sign_photos, load_day_records_local
from image_utils import resize_image_fixed
from xml_utils import inject_images_into_docx, ensure_dir
from db_utils import save_download_db, load_download_db


# ---------------------- helper utils ----------------------

def inline_replace_paragraph(paragraph, target, replacement):
    if target not in paragraph.text:
        return False
    new_text = paragraph.text.replace(target, str(replacement))
    for r in paragraph.runs:
        r.text = ""
    paragraph.add_run(new_text)
    return True


def replace_text_placeholders(doc, mapping):
    for p in doc.paragraphs:
        for key, val in mapping.items():
            if key in p.text:
                inline_replace_paragraph(p, key, val)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in mapping.items():
                        if key in p.text:
                            inline_replace_paragraph(p, key, val)

    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                for key, val in mapping.items():
                    if key in p.text:
                        inline_replace_paragraph(p, key, val)
            for p in section.footer.paragraphs:
                for key, val in mapping.items():
                    if key in p.text:
                        inline_replace_paragraph(p, key, val)
    except Exception:
        pass


def insert_image_at_placeholder(doc, placeholder, image_path, width_inches=2.8):
    """
    Fallback insertion using python-docx when placeholder is plain text (not inside drawing/textbox).
    Returns number of insertions done.
    """
    width = Inches(width_inches)
    inserted = 0

    def clean_text(s):
        if s is None:
            return ""
        return s.replace(" ", "").replace("\t", "").replace("\n", "").replace("\r", "")

    def clear_and_insert(paragraph, cell=None):
        nonlocal inserted
        try:
            if cell:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = ""
                p = cell.paragraphs[0]
            else:
                for r in paragraph.runs:
                    r.text = ""
                p = paragraph
            run = p.add_run()
            run.add_picture(image_path, width=width)
            inserted += 1
        except Exception as e:
            log(f"Image insert failed at {placeholder}: {e}")

    for p in doc.paragraphs:
        if placeholder in clean_text(p.text):
            clear_and_insert(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in clean_text(p.text):
                        clear_and_insert(p, cell)

    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                if placeholder in clean_text(p.text):
                    clear_and_insert(p)
            for p in section.footer.paragraphs:
                if placeholder in clean_text(p.text):
                    clear_and_insert(p)
    except Exception:
        pass

    return inserted


def force_arial(doc, size_pt=11):
    for p in doc.paragraphs:
        for run in p.runs:
            try:
                run.font.name = "Arial"
                run.font.size = Pt(size_pt)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            except Exception:
                pass
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        try:
                            run.font.name = "Arial"
                            run.font.size = Pt(size_pt)
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                        except Exception:
                            pass


def safe_save_docx(base_path):
    """
    If file exists and is locked by Word, returns a new versioned filename
    (e.g., base_v2.docx) so we don't raise PermissionError when writing.
    """
    if not os.path.exists(base_path):
        return base_path
    try:
        os.rename(base_path, base_path)
        return base_path
    except PermissionError:
        folder = os.path.dirname(base_path)
        name = os.path.basename(base_path)
        base, ext = os.path.splitext(name)
        version = 2
        while True:
            new_name = f"{base}_v{version}{ext}"
            new_path = os.path.join(folder, new_name)
            if not os.path.exists(new_path):
                return new_path
            version += 1


def find_placeholders_in_docx_xml(docx_path, placeholders):
    """
    Return dict placeholder -> count of raw occurrences in main document.xml + common headers/footers.
    This checks the raw XML text for exact literal occurrences and some whitespace/encoding variants.
    """
    counts = {p: 0 for p in placeholders}
    try:
        with zipfile.ZipFile(docx_path, 'r') as z:
            parts = [
                'word/document.xml',
                'word/header1.xml', 'word/footer1.xml',
                'word/header2.xml', 'word/footer2.xml',
                'word/header3.xml', 'word/footer3.xml'
            ]
            for part in parts:
                try:
                    data = z.read(part).decode('utf-8', errors='ignore')
                except KeyError:
                    continue
                for p in placeholders:
                    # literal occurrences
                    counts[p] += data.count(p)
                    # regex variant to catch whitespace/NBSP/zero-width splits inside parentheses
                    try:
                        inner = re.escape(p.strip("()"))
                        pattern = r"\(" + r"[\s\u00A0\u200B\u00AD]*" + inner + r"[\s\u00A0\u200B\u00AD]*" + r"\)"
                        matches = re.findall(pattern, data)
                        counts[p] += len(matches)
                    except re.error:
                        pass
    except Exception as e:
        log(f"find_placeholders_in_docx_xml: failed to inspect {docx_path}: {e}")
    return counts


# ---------------------- photo / placeholder builders ----------------------

def build_pic_placeholders_map(date_str, desired_w=162, desired_h=162):
    """
    Scans LOCAL_DIR/<date>/photos and returns dict mapping '(pic_<cage>)' -> resized_image_path
    Example: '002.jpg' -> placeholder '(pic_002)' and '(pic_2)' (both)
    """
    photos_dir = os.path.join(LOCAL_DIR, date_str, "photos")
    mapping = {}
    if not os.path.exists(photos_dir):
        return mapping

    for fname in sorted(os.listdir(photos_dir)):
        if not fname.lower().endswith((".jpg", ".jpeg", ".png")):
            continue
        src = os.path.join(photos_dir, fname)

        name_root = os.path.splitext(fname)[0]
        digits = ''.join([c for c in name_root if c.isdigit()])
        if not digits:
            digits = name_root

        placeholder1 = f"(pic_{digits})"
        placeholder2 = f"(pic_{int(digits)})" if digits.isdigit() else placeholder1

        base, ext = os.path.splitext(src)
        resized_path = base + "_162.jpg"
        try:
            resize_image_fixed(src, resized_path, desired_w, desired_h)
            mapping[placeholder1] = resized_path
            if placeholder2 != placeholder1:
                mapping[placeholder2] = resized_path
            log(f"Prepared resized cage photo {fname} -> {os.path.basename(resized_path)}")
        except Exception as e:
            log(f"Resize failed for {src}: {e}")
            mapping[placeholder1] = src
            if placeholder2 != placeholder1:
                mapping[placeholder2] = src

    return mapping


# ---------------------- places & totals config ----------------------

PLACES_CAGES = {
    "Southern Promenade": [458,459,460,461,462,463,464,465,466,467,468,469,470,471,472,473,474],
    "Eastern Promenade": [475,476,477,478,526,527,528,530,531],
    "U-shape East & West Wing": [484,485,486,487,488,489,490,491,492,501,502,505],
    "Marina Carpark 2A": [506,507,508,510],
    "Marina Carpark 2B": [493,494,495],
    "Northern Promenade": [512,513],
    "QD Complex (External)": [496,497,498,499,500,504,509],
    "Crescent Park 01": [523,524,525,540,541,542,543,544,545,546,547],
    "Crescent Park 02": [479,480,481,482,483],
    "Crescent Park 03": [514,516,517,518,519,520,521,522],
    "Crescent Park 04": [503,511,532,533,534,535,536,537,538,539],
    "Crescent Park 05": [549,550,551,552,553,554,555,556,557,558,559,560,561,562,563,564],
    "Al Khuzama Zone -2": [604,605,606,607,608,609,610,611,612,613,614,615,617,618,619,620],
    "Al Khuzama Zone -1": [588,589,590,591,592,593,594,595,596,597,598,599,600,601,602,603],
    "Al Nafel Park": [577,578,579,580,581],
    "QETAIFAN ZONE 1": [569,570,574,575],
    "QETAIFAN ZONE 2": [567,572,573,576],
    "QETAIFAN ZONE 3": [565,566,568],
    "Qetaifan North Park": [623,624,625,626,630,631,632,633,634],
    "Road A1 - Al Khuzama": [621,622,627,628,629],
    "Seef Lusail North": [635,636,637,638,639,640,641,642],
}

PLACE_TOTAL_PLACEHOLDERS = {
    # SHIFT 1
    "Southern Promenade_1": ("(1spt)", "(1spm)", "(1spl)"),
    "Eastern Promenade_1": ("(1ept)", "(1epm)", "(1epl)"),
    "U-shape East & West Wing_1": ("(1uset)", "(1usem)", "(1usel)"),
    "Northern Promenade_1": ("(1npt)", "(1npm)", "(1npl)"),
    "QD Complex (External)_1": ("(1qdcet)", "(1qdcem)", "(1qdcel)"),
    "Marina Carpark 2A_1": ("(1mc2at)", "(1mc2am)", "(1mc2al)"),
    "Marina Carpark 2B_1": ("(1mc2bt)", "(1mc2bm)", "(1mc2bl)"),
    "Crescent Park 01_1": ("(1cp1t)", "(1cp1m)", "(1cp1l)"),
    "Crescent Park 02_1": ("(1cp2t)", "(1cp2m)", "(1cp2l)"),
    "Crescent Park 03_1": ("(1cp3t)", "(1cp3m)", "(1cp3l)"),
    "Crescent Park 04_1": ("(1cp4t)", "(1cp4m)", "(1cp4l)"),
    "Crescent Park 05_1": ("(1cp5t)", "(1cp5m)", "(1cp5l)"),
    "QETAIFAN ZONE 1_1": ("(1qz1t)", "(1qz1m)", "(1qz1l)"),
    "QETAIFAN ZONE 2_1": ("(1qz2t)", "(1qz2m)", "(1qz2l)"),
    "QETAIFAN ZONE 3_1": ("(1qz3t)", "(1qz3m)", "(1qz3l)"),
    "Al Nafel Park_1": ("(1anpt)", "(1anpm)", "(1anpl)"),
    "Al Khuzama Zone -2_1": ("(1akz2t)", "(1akz2m)", "(1akz2l)"),
    "Al Khuzama Zone -1_1": ("(1akz1t)", "(1akz1m)", "(1akz1l)"),
    "Road A1 - Al Khuzama_1": ("(1raakt)", "(1raakm)", "(1raakl)"),
    "Qetaifan North Park_1": ("(1qnpt)", "(1qnpm)", "(1qnpl)"),
    "Seef Lusail North_1": ("(1slnt)", "(1slnm)", "(1slnl)"),

    # SHIFT 2
    "Southern Promenade_2": ("(2spt)", "(2spm)", "(2spl)"),
    "Eastern Promenade_2": ("(2ept)", "(2epm)", "(2epl)"),
    "U-shape East & West Wing_2": ("(2uset)", "(2usem)", "(2usel)"),
    "Northern Promenade_2": ("(2npt)", "(2npm)", "(2npl)"),
    "QD Complex (External)_2": ("(2qdcet)", "(2qdcem)", "(2qdcel)"),
    "Marina Carpark 2A_2": ("(2mc2at)", "(2mc2am)", "(2mc2al)"),
    "Marina Carpark 2B_2": ("(2mc2bt)", "(2mc2bm)", "(2mc2bl)"),
    "Crescent Park 01_2": ("(2cp1t)", "(2cp1m)", "(2cp1l)"),
    "Crescent Park 02_2": ("(2cp2t)", "(2cp2m)", "(2cp2l)"),
    "Crescent Park 03_2": ("(2cp3t)", "(2cp3m)", "(2cp3l)"),
    "Crescent Park 04_2": ("(2cp4t)", "(2cp4m)", "(2cp4l)"),
    "Crescent Park 05_2": ("(2cp5t)", "(2cp5m)", "(2cp5l)"),
    "QETAIFAN ZONE 1_2": ("(2qz1t)", "(2qz1m)", "(2qz1l)"),
    "QETAIFAN ZONE 2_2": ("(2qz2t)", "(2qz2m)", "(2qz2l)"),
    "QETAIFAN ZONE 3_2": ("(2qz3t)", "(2qz3m)", "(2qz3l)"),
    "Al Nafel Park_2": ("(2anpt)", "(2anpm)", "(2anpl)"),
    "Al Khuzama Zone -2_2": ("(2akz2t)", "(2akz2m)", "(2akz2l)"),
    "Al Khuzama Zone -1_2": ("(2akz1t)", "(2akz1m)", "(2akz1l)"),
    "Road A1 - Al Khuzama_2": ("(2raakt)", "(2raakm)", "(2raakl)"),
    "Qetaifan North Park_2": ("(2qnpt)", "(2qnpm)", "(2qnpl)"),
    "Seef Lusail North_2": ("(2slnt)", "(2slnm)", "(2slnl)")
}

# PLACE_TOTAL_PLACEHOLDERS = {
#     # SHIFT 1
#     "Southern Promenade_1": ("(1spt)", "(1spm)", "(1spl)"),
#     "Eastern Promenade_1": ("(1ept)", "(1epm)", "(1epl)"),
#     "U-shape East & West Wing_1": ("(1uset)", "(1usem)", "(1usel)"),
#     "Northern Promenade_1": ("(1npt)", "(1npm)", "(1npl)"),
#     "QD Complex (External)_1": ("(1qdcet)", "(1qdcem)", "(1qdcel)"),
#     "Marina Carpark 2A_1": ("(1mc2at)", "(1mc2am)", "(1mc2al)"),
#     "Marina Carpark 2B_1": ("(1mc2bt)", "(1mc2bm)", "(1mc2bl)"),
#     "Crescent Park 01_1": ("(1cp1t)", "(1cp1m)", "(1cp1l)"),
#     "Crescent Park 02_1": ("(1cp2t)", "(1cp2m)", "(1cp2l)"),
#     "Crescent Park 03_1": ("(1cp3t)", "(1cp3m)", "(1cp3l)"),
#     "Crescent Park 04_1": ("(1cp4t)", "(1cp4m)", "(1cp4l)"),
#     "Crescent Park 05_1": ("(1cp5t)", "(1cp5m)", "(1cp5l)"),
#     "QATAIFAN ZONE 1_1": ("(1qz1t)", "(1qz1m)", "(1qz1l)"),
#     "QATAIFAN ZONE 2_1": ("(1qz2t)", "(1qz2m)", "(1qz2l)"),
#     "QATAIFAN ZONE 3_1": ("(1qz3t)", "(1qz3m)", "(1qz3l)"),
#     "Al Nafel Park_1": ("(1anpt)", "(1anpm)", "(1anpl)"),
#     "Al Khuzama Zone -2_1": ("(1akz2t)", "(1akz2m)", "(1akz2l)"),
#     "Al Khuzama Zone -1_1": ("(1akz1t)", "(1akz1m)", "(1akz1l)"),
#     "Road A1 - Al Khuzama_1": ("(1raakt)", "(1raakm)", "(1raakl)"),
#     "Qetaifan North Park_1": ("(1qnpt)", "(1qnpm)", "(1qnpl)"),
#     "Seef Lusail North_1": ("(1slnt)", "(1slnm)", "(1slnl)"),

#     # SHIFT 2
#     "Southern Promenade_2": ("(2spt)", "(2spm)", "(2spl)"),
#     "Eastern Promenade_2": ("(2ept)", "(2epm)", "(2epl)"),
#     "U-shape East & West Wing_2": ("(2uset)", "(2usem)", "(2usel)"),
#     "Northern Promenade_2": ("(2npt)", "(2npm)", "(2npl)"),
#     "QD Complex (External)_2": ("(2qdcet)", "(2qdcem)", "(2qdcel)"),
#     "Marina Carpark 2A_2": ("(2mc2at)", "(2mc2am)", "(2mc2al)"),
#     "Marina Carpark 2B_2": ("(2mc2bt)", "(2mc2bm)", "(2mc2bl)"),
#     "Crescent Park 01_2": ("(2cp1t)", "(2cp1m)", "(2cp1l)"),
#     "Crescent Park 02_2": ("(2cp2t)", "(2cp2m)", "(2cp2l)"),
#     "Crescent Park 03_2": ("(2cp3t)", "(2cp3m)", "(2cp3l)"),
#     "Crescent Park 04_2": ("(2cp4t)", "(2cp4m)", "(2cp4l)"),
#     "Crescent Park 05_2": ("(2cp5t)", "(2cp5m)", "(2cp5l)"),
#     "QATAIFAN ZONE 1_2": ("(2qz1t)", "(2qz1m)", "(2qz1l)"),
#     "QATAIFAN ZONE 2_2": ("(2qz2t)", "(2qz2m)", "(2qz2l)"),
#     "QATAIFAN ZONE 3_2": ("(2qz3t)", "(2qz3m)", "(2qz3l)"),
#     "Al Nafel Park_2": ("(2anpt)", "(2anpm)", "(2anpl)"),
#     "Al Khuzama Zone -2_2": ("(2akz2t)", "(2akz2m)", "(2akz2l)"),
#     "Al Khuzama Zone -1_2": ("(2akz1t)", "(2akz1m)", "(2akz1l)"),
#     "Road A1 - Al Khuzama_2": ("(2raakt)", "(2raakm)", "(2raakl)"),
#     "Qetaifan North Park_2": ("(2qnpt)", "(2qnpm)", "(2qnpl)"),
#     "Seef Lusail North_2": ("(2slnt)", "(2slnm)", "(2slnl)")
# }


def cage_placeholder(shift, cage_number):
    return f"{shift}c{cage_number}"


def pic_placeholder(cage_number):
    return f"(pic_{cage_number})"


# ---------------------- record processing ----------------------

def process_record_updates(date_str, local_dir, resize_fn, logger):
    """
    Reads local JSON records for date_str and builds:
      - text_map: mapping placeholder -> string value (per-cage counts and totals)
      - pic_map: mapping pic placeholders like (pic_613) -> resized image path
    """
    text_map = {}
    pic_map = {}

    data_folder = os.path.join(local_dir, date_str, "data")
    photos_folder = os.path.join(local_dir, date_str, "photos")

    records = []
    if os.path.exists(data_folder):
        for fname in sorted(os.listdir(data_folder)):
            if not fname.lower().endswith(".json"):
                continue
            try:
                with open(os.path.join(data_folder, fname), "r", encoding="utf-8") as f:
                    r = json.load(f)
                    r["_filename"] = fname
                    records.append(r)
            except Exception as e:
                logger(f"process_record_updates: failed reading {fname}: {e}")

    # initialize placeholders for both shifts (per-cage)
    for place, cages in PLACES_CAGES.items():
        for c in cages:
            text_map[f"(1c{c})"] = "0"
            text_map[f"(2c{c})"] = "0"

    # initialize totals placeholders to zero so totals always exist (prevents missing replacements)
    for key, tpl in PLACE_TOTAL_PLACEHOLDERS.items():
        tot_ph, my_ph, loc_ph = tpl
        text_map[tot_ph] = "0"
        text_map[my_ph] = "0"
        text_map[loc_ph] = "0"

    # aggregate structure per place
    place_agg = {
        place: {
            "myna_shift1": 0, "local_shift1": 0, "total_shift1": 0,
            "myna_shift2": 0, "local_shift2": 0, "total_shift2": 0
        }
        for place in PLACES_CAGES.keys()
    }

    # process all record_update JSON files
    for r in records:
        if r.get("type") != "record_update":
            continue

        logger(f"[DEBUG] Loaded JSON record: {json.dumps(r, indent=2)}")

        shift = str(r.get("shift", "1")).strip()

        try:
            cage_no = int(r.get("cage_number"))
        except:
            logger("[DEBUG] Invalid cage number")
            continue

        logger(f"[DEBUG] Extracted cage number: {cage_no} for shift {shift}")

        try:
            myna = int(str(r.get("myna_captured") or "0"))
        except Exception:
            myna = 0
        try:
            local = int(str(r.get("local_released") or "0"))
        except Exception:
            local = 0

        total = myna + local

        ph = f"({shift}c{cage_no})"
        logger(f"[DEBUG] Placeholder for this cage: {ph}")

        if ph in text_map:
            logger(f"[DEBUG] Placeholder FOUND in text_map: {ph}")
        else:
            logger(f"[DEBUG] Placeholder NOT FOUND in text_map: {ph}")

        text_map[ph] = f"{myna}M,{local}L"
        logger(f"[DEBUG] Updated placeholder {ph} -> {myna}M,{local}L")

        opp_shift = "1" if shift == "2" else "2"
        opp_ph = f"({opp_shift}c{cage_no})"
        text_map[opp_ph] = "0"
        logger(f"[DEBUG] Zeroed opposite shift placeholder: {opp_ph}")

        # find the place for this cage and aggregate
        for place, cages in PLACES_CAGES.items():
            if cage_no in cages:
                if shift == "1":
                    place_agg[place]["myna_shift1"] += myna
                    place_agg[place]["local_shift1"] += local
                    place_agg[place]["total_shift1"] += total
                else:
                    place_agg[place]["myna_shift2"] += myna
                    place_agg[place]["local_shift2"] += local
                    place_agg[place]["total_shift2"] += total
                break

        # picture processing
        photo = r.get("photo")
        if photo:
            src = os.path.join(photos_folder, photo)
            if os.path.exists(src):
                base, ext = os.path.splitext(src)
                dst = base + "_162.jpg"
                try:
                    resize_fn(src, dst, 162, 162)
                    pic_map[f"(pic_{cage_no})"] = dst
                except Exception as e:
                    logger(f"process_record_updates: resize failed for {src}: {e}")
                    pic_map[f"(pic_{cage_no})"] = src

    # now fill totals placeholders for *both shifts*
    for place, agg in place_agg.items():
        for shift in ("1", "2"):
            key = f"{place}_{shift}"
            if key not in PLACE_TOTAL_PLACEHOLDERS:
                continue

            total_ph, myna_ph, local_ph = PLACE_TOTAL_PLACEHOLDERS[key]

            if shift == "1":
                t = agg["total_shift1"]
                m = agg["myna_shift1"]
                l = agg["local_shift1"]
            else:
                t = agg["total_shift2"]
                m = agg["myna_shift2"]
                l = agg["local_shift2"]

            text_map[total_ph] = str(t)
            text_map[myna_ph] = str(m)
            text_map[local_ph] = str(l)

    logger(f"process_record_updates: generated {len(text_map)} text placeholders and {len(pic_map)} pic placeholders")
    return text_map, pic_map


# ---------------------- main orchestration ----------------------

def create_partial_report_with_shift_signs(date_str):
    log(f"Creating partial report for {date_str}")

    # find sign photos
    sign_map = find_shift_sign_photos(date_str)

    # Load original template
    try:
        doc = Document(TEMPLATE_ORIG)
    except Exception as e:
        log(f"Failed to open template {TEMPLATE_ORIG}: {e}")
        return None

    # Date strings
    human_date = datetime.strptime(date_str, "%Y-%m-%d").strftime("%d %B %Y")

    # gather updates first (text + record images)
    try:
        text_map_updates, pic_map_from_records = process_record_updates(
            date_str, LOCAL_DIR, resize_image_fixed, log
        )
        text_map_updates = text_map_updates or {}
        pic_map_from_records = pic_map_from_records or {}
        log(f"Received {len(text_map_updates)} text updates and {len(pic_map_from_records)} record images")
    except Exception as e:
        log(f"process_record_updates failed: {e}")
        text_map_updates = {}
        pic_map_from_records = {}

    # Build placeholders -> images map for sign placeholders
    placeholder_image_map = {
        "(shift_1_signin)": None,
        "(shift_1_signout)": None,
        "(shift_2_signin)": None,
        "(shift_2_signout)": None
    }

    for k in ["shift_1_signin", "shift_1_signout", "shift_2_signin", "shift_2_signout"]:
        src = sign_map.get(k)
        if src and os.path.exists(src):
            base, ext = os.path.splitext(src)
            resized = base + "_162.jpg"
            try:
                resize_image_fixed(src, resized, 162, 162)
                placeholder_image_map[f"({k})"] = resized
                log(f"Resized sign {k}: {resized}")
            except Exception as e:
                log(f"Sign resize failed for {src}: {e}")
                placeholder_image_map[f"({k})"] = src
        else:
            placeholder_image_map[f"({k})"] = None

    # build general pic_map (scanning photos folder)
    try:
        pic_map = build_pic_placeholders_map(date_str, desired_w=162, desired_h=162)
    except Exception as e:
        log(f"build_pic_placeholders_map failed: {e}")
        pic_map = {}

    # merged_pic_map: scanned pics overridden by any record-specific pics
    merged_pic_map = {}
    merged_pic_map.update(pic_map)
    merged_pic_map.update(pic_map_from_records or {})

    # integrate into placeholder_image_map
    placeholder_image_map.update(merged_pic_map)

    # mapping_for_xml contains date and all numeric placeholders + totals
    mapping_for_xml = {
        "(date)": date_str,
        "(date_with_month)": human_date
    }

    # merge numeric placeholders and totals (these are parenthesized keys)
    mapping_for_xml.update(text_map_updates)

    # ensure output dir exists, save a tmp docx where python-docx does visible replacements
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    tmp_text_docx = os.path.join(OUTPUT_DIR, f"temp_text_{date_str}_{uuid.uuid4().hex}.docx")
    try:
        replace_text_placeholders(doc, mapping_for_xml)
        doc.save(tmp_text_docx)
        log(f"Applied {len(mapping_for_xml)} visible text replacements via python-docx")
    except Exception as e:
        log(f"Failed saving visible-text update docx: {e}")
        return None

    # --- Debug: inspect saved tmp_text_docx for placeholders actually present in XML ---
    try:
        expected_phs = [k for k in mapping_for_xml.keys() if k.startswith("(") and k.endswith(")")]
        expected_phs = sorted(set(expected_phs))
        xml_counts = find_placeholders_in_docx_xml(tmp_text_docx, expected_phs)
        missing = [p for p, c in xml_counts.items() if c == 0]
        present = [p for p, c in xml_counts.items() if c > 0]
        log(f"XML placeholder scan: {len(present)} found, {len(missing)} missing (checked {len(expected_phs)})")
        if missing:
            sample = missing[:30]
            log(f"Missing placeholders (sample up to 30): {sample}")
    except Exception as e:
        log(f"XML placeholder scan failed: {e}")

    # final doc path (safe)
    final_docx = os.path.join(OUTPUT_DIR, f"Daily_Report_{date_str}_partial.docx")
    final_docx_safe = safe_save_docx(final_docx)

    # fill mapping_for_xml also with non-parenthesized variants (some injector templates expect both)
    for k, v in text_map_updates.items():
        mapping_for_xml[k] = v
        if k.startswith("(") and k.endswith(")"):
            mapping_for_xml[k[1:-1]] = v

    log(f"XML text replacements prepared: {len(mapping_for_xml)} entries (includes non-parenthesized variants)")

    # Perform XML-level injection (images + text_map_for_xml)
    try:
        inject_images_into_docx(tmp_text_docx, final_docx_safe, placeholder_image_map, text_map=mapping_for_xml)
    except Exception as e:
        log(f"XML injection failed: {e}")
        # fallback copy
        try:
            shutil.copy2(tmp_text_docx, final_docx_safe)
        except Exception as e2:
            log(f"Failed fallback copy: {e2}")
            return None

    # Final python-docx pass to force Arial and do any plain-text picture insertions fallback
    try:
        doc_final = Document(final_docx_safe)
        # fallback: if some pic placeholders weren't handled at XML level and are plain text, try python-docx insertion
        for placeholder, img_path in placeholder_image_map.items():
            if img_path and "(" in placeholder:
                inserted = insert_image_at_placeholder(doc_final, placeholder, img_path, width_inches=1.8)
                if inserted:
                    log(f"Fallback inserted {placeholder} via python-docx (count={inserted})")
        force_arial(doc_final)
        doc_final.save(final_docx_safe)
    except Exception as e:
        log(f"Post python-docx formatting failed: {e}")

    # cleanup temp file
    try:
        if os.path.exists(tmp_text_docx):
            os.remove(tmp_text_docx)
    except Exception:
        pass

    log(f"Saved partial: {final_docx_safe}")
    return final_docx_safe
