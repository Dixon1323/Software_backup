from docx import Document

TEMPLATE_PATH = r"D:\Dixon\Automation_software\Office Software\template_formatted.docx"

doc = Document(TEMPLATE_PATH)

print("\n=== PARAGRAPHS ===")
for i, p in enumerate(doc.paragraphs):
    print(i, "=>", repr(p.text))

print("\n=== TABLE CELLS ===")
for tid, table in enumerate(doc.tables):
    for rid, row in enumerate(table.rows):
        for cid, cell in enumerate(row.cells):
            cell_text = " | ".join([p.text for p in cell.paragraphs])
            print(f"Table {tid} Row {rid} Col {cid} => {repr(cell_text)}")
