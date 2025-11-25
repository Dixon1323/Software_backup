import requests
import os
import time
import json
import zipfile
import shutil
import uuid
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from PIL import Image
from lxml import etree

# -------------------------
# CONFIG
# -------------------------
BASE_URL = "https://birdportal.pythonanywhere.com/records/"
LOCAL_DIR = "D:/Dixon/Automation_software/Office Software/sync/records/"

DOWNLOADED_DB = "D:/Dixon/Automation_software/Office Software/sync/downloaded_files.json"
LOG_FILE = "D:/Dixon/Automation_software/Office Software/sync/sync.log"

LOOP_INTERVAL = 10  # seconds

# original template uploaded into the environment
TEMPLATE_ORIG = r"D:\Dixon\Automation_software\Office Software\template_formatted.docx"
OUTPUT_DIR = "D:/Dixon/Automation_software/Office Software/sync/reports/"

# Image media extension to save to docx media folder
MEDIA_EXT = ".png"
# EMU scale used for pixel -> EMU conversion
EMU_PER_PIXEL = 9525

# namespace map for building drawing snippet
NSMAP = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
    'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
}

# -------------------------
# UTIL: logging
# -------------------------
def log(msg):
    timestamp = datetime.now().strftime("[%Y-%m-%d %H:%M:%S]")
    line = f"{timestamp} {msg}"
    print(line)
    try:
        with open(LOG_FILE, "a", encoding="utf-8") as f:
            f.write(line + "\n")
    except Exception:
        pass

# -------------------------
# HTTP helpers
# -------------------------
def safe_request(url, retries=3):
    for i in range(retries):
        try:
            res = requests.get(url, timeout=10)
            if res.status_code == 200:
                return res
            else:
                log(f"Bad response {res.status_code}: {url}")
        except Exception as e:
            log(f"Network error accessing {url} ({i+1}/{retries}): {e}")
        time.sleep(2)
    return None


# -------------------------
# DOWNLOAD FILE
# -------------------------
def download_file(url, local_path):
    res = safe_request(url)
    if res is None:
        log(f"FAILED downloading: {url}")
        return False
    os.makedirs(os.path.dirname(local_path), exist_ok=True)
    try:
        with open(local_path, "wb") as f:
            for chunk in res.iter_content(1024):
                f.write(chunk)
        log(f"Downloaded: {local_path}")
        return True
    except Exception as e:
        log(f"Failed saving download {local_path}: {e}")
        return False
    

# -------------------------
# DELETE FROM SERVER
# -------------------------
def delete_from_server(day):
    url = BASE_URL + f"{day}/delete"
    try:
        res = requests.post(url)
        if res.status_code == 200:
            log(f"Server files deleted for {day}")
        else:
            log(f"Failed to delete server files for {day}: HTTP {res.status_code}")
    except Exception as e:
        log(f"Error deleting files on server for {day}: {e}")

# -------------------------
# local DB (download tracking)
# -------------------------
def load_download_db():
    if not os.path.exists(DOWNLOADED_DB):
        return {}
    try:
        return json.load(open(DOWNLOADED_DB, encoding="utf-8"))
    except Exception:
        return {}

def save_download_db(db):
    with open(DOWNLOADED_DB, "w", encoding="utf-8") as f:
        json.dump(db, f, indent=4)

download_db = load_download_db()

# -------------------------
# RESIZE IMAGE
# -------------------------
def resize_image_fixed(input_path, output_path, width, height):
    """Resize to fixed width/height and ensure JPEG-compatible RGB mode."""
    try:
        img = Image.open(input_path)

        # Convert RGBA → RGB (JPEG does NOT support transparency)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")

        img = img.resize((width, height), Image.LANCZOS)
        # Ensure parent exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        # Save as JPEG (we will re-open and convert to PNG when injecting if needed)
        img.save(output_path, format="JPEG")
        img.close()

    except Exception as e:
        raise RuntimeError(f"Resize error for {input_path}: {e}")

    

def safe_save_docx(base_path):
    """
    Prevent PermissionError when DOCX is open.
    If locked, save as base_path_v2.docx, v3... etc.
    """
    if not os.path.exists(base_path):
        return base_path  # safe

    # test if file is locked by trying rename-to-self
    try:
        os.rename(base_path, base_path)
        return base_path  # not locked
    except PermissionError:
        # file is locked → create next version
        folder = os.path.dirname(base_path)
        name = os.path.basename(base_path)
        base, ext = os.path.splitext(name)

        version = 2
        while True:
            new_name = f"{base}_v{version}{ext}"
            new_path = os.path.join(folder, new_name)
            if not os.path.exists(new_path):
                return new_path
            version += 1


# -------------------------
# Placeholder text replacement utilities (python-docx friendly)
# -------------------------
def inline_replace_paragraph(paragraph, target, replacement):
    if target not in paragraph.text:
        return False
    new_text = paragraph.text.replace(target, str(replacement))
    # Clear runs
    for r in paragraph.runs:
        r.text = ""
    paragraph.add_run(new_text)
    return True

def replace_text_placeholders(doc, mapping):
    # paragraphs
    for p in doc.paragraphs:
        for key, val in mapping.items():
            if key in p.text:
                inline_replace_paragraph(p, key, val)
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, val in mapping.items():
                        if key in p.text:
                            inline_replace_paragraph(p, key, val)
    # headers/footers
    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                for key, val in mapping.items():
                    if key in p.text:
                        inline_replace_paragraph(p, key, val)
            for p in section.footer.paragraphs:
                for key, val in mapping.items():
                    if key in p.text:
                        inline_replace_paragraph(p, key, val)
    except Exception:
        pass

# -------------------------
# helpers: force Arial formatting
# -------------------------
def force_arial(doc, size_pt=11):
    for p in doc.paragraphs:
        for run in p.runs:
            try:
                run.font.name = "Arial"
                run.font.size = Pt(size_pt)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
            except Exception:
                pass
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        try:
                            run.font.name = "Arial"
                            run.font.size = Pt(size_pt)
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
                        except Exception:
                            pass

# -------------------------
# load local JSON records for a given date
# -------------------------
def load_day_records_local(date_str):
    data_folder = os.path.join(LOCAL_DIR, date_str, "data")
    records = []
    if not os.path.exists(data_folder):
        return records
    for fname in sorted(os.listdir(data_folder)):
        if not fname.lower().endswith(".json"):
            continue
        try:
            with open(os.path.join(data_folder, fname), "r", encoding="utf-8") as f:
                rec = json.load(f)
                rec["_filename"] = fname
                records.append(rec)
        except Exception as e:
            log(f"Could not read {fname}: {e}")
    return records

# -------------------------
# find shift sign images from local synced photos
# -------------------------
def find_shift_sign_photos(date_str):
    result = {
        "shift_1_signin": None,
        "shift_1_signout": None,
        "shift_2_signin": None,
        "shift_2_signout": None
    }
    records = load_day_records_local(date_str)
    photo_dir = os.path.join(LOCAL_DIR, date_str, "photos")
    for r in records:
        r_type = r.get("type")
        shift = str(r.get("shift", ""))
        photo = r.get("photo")
        if not photo:
            continue
        full_path = os.path.join(photo_dir, photo)
        if not os.path.exists(full_path):
            continue
        # pick first matching photo for each type
        if r_type == "start_shift" and shift == "1" and result["shift_1_signin"] is None:
            result["shift_1_signin"] = full_path
        if r_type == "end_shift" and shift == "1" and result["shift_1_signout"] is None:
            result["shift_1_signout"] = full_path
        if r_type == "start_shift" and shift == "2" and result["shift_2_signin"] is None:
            result["shift_2_signin"] = full_path
        if r_type == "end_shift" and shift == "2" and result["shift_2_signout"] is None:
            result["shift_2_signout"] = full_path
    return result

# -------------------------
# Insert image into python-docx placeholders (fallback)
# (still used for placeholders that python-docx can see as normal text)
# -------------------------
def insert_image_at_placeholder(doc, placeholder, image_path, width_inches=2.8):
    width = Inches(width_inches)
    inserted = 0

    def clean_text(s):
        if s is None:
            return ""
        return s.replace(" ", "").replace("\t", "").replace("\n", "").replace("\r", "")

    def clear_and_insert(paragraph, cell=None):
        nonlocal inserted
        try:
            if cell:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.text = ""
                p = cell.paragraphs[0]
            else:
                for r in paragraph.runs:
                    r.text = ""
                p = paragraph
            run = p.add_run()
            run.add_picture(image_path, width=width)
            inserted += 1
        except Exception as e:
            log(f"Image insert failed at {placeholder}: {e}")

    for p in doc.paragraphs:
        if placeholder in clean_text(p.text):
            clear_and_insert(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in clean_text(p.text):
                        clear_and_insert(p, cell)

    try:
        for section in doc.sections:
            for p in section.header.paragraphs:
                if placeholder in clean_text(p.text):
                    clear_and_insert(p)
            for p in section.footer.paragraphs:
                if placeholder in clean_text(p.text):
                    clear_and_insert(p)
    except Exception:
        pass

    return inserted

# -------------------------
# XML image injection helpers (inject images into docx XML where placeholders are)
# -------------------------
def ensure_dir(path):
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)

def next_media_name(media_dir):
    existing = set(os.listdir(media_dir)) if os.path.exists(media_dir) else set()
    i = 1
    while True:
        name = f"image{i:03d}{MEDIA_EXT}"
        if name not in existing:
            return name
        i += 1

def add_image_file_to_media(tmpdir, image_src):
    media_dir = os.path.join(tmpdir, "word", "media")
    ensure_dir(media_dir)
    fname = next_media_name(media_dir)
    out_path = os.path.join(media_dir, fname)
    img = Image.open(image_src)
    # save as chosen format
    try:
        if MEDIA_EXT.lower() == ".png":
            # convert/save into PNG media file
            img.save(out_path, format="PNG")
        else:
            # save as-is (PIL will pick format based on ext)
            img.save(out_path)
        w_px, h_px = img.size
        img.close()
        return fname, w_px, h_px
    except Exception as e:
        img.close()
        raise

def ensure_rels_file(rels_full):
    if not os.path.exists(rels_full):
        # create minimal Relationships xml root
        root = etree.Element("{http://schemas.openxmlformats.org/package/2006/relationships}Relationships")
        tree = etree.ElementTree(root)
        tree.write(rels_full, xml_declaration=True, encoding="utf-8")

def add_image_relationship(rels_full_path, target_media_filename):
    parser = etree.XMLParser(remove_blank_text=True)
    tree = etree.parse(rels_full_path, parser)
    root = tree.getroot()
    existing_ids = [el.get('Id') for el in root.findall('.//{http://schemas.openxmlformats.org/package/2006/relationships}Relationship')]
    maxn = 0
    for eid in existing_ids:
        if eid and eid.startswith("rId"):
            try:
                n = int(eid[3:])
                if n > maxn:
                    maxn = n
            except Exception:
                pass
    new_id = f"rId{maxn + 1}"
    RelTag = "{http://schemas.openxmlformats.org/package/2006/relationships}Relationship"
    rel = etree.SubElement(root, RelTag)
    rel.set("Id", new_id)
    rel.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/image")
    rel.set("Target", "media/" + target_media_filename)
    tree.write(rels_full_path, xml_declaration=True, encoding="utf-8")
    return new_id

def build_drawing_xml(rel_id, cx, cy):
    drawing_xml = f'''
    <w:r xmlns:w="{NSMAP['w']}" xmlns:wp="{NSMAP['wp']}" xmlns:a="{NSMAP['a']}" xmlns:pic="{NSMAP['pic']}" xmlns:r="{NSMAP['r']}">
      <w:drawing>
        <wp:inline distT="0" distB="0" distL="0" distR="0">
          <wp:extent cx="{cx}" cy="{cy}"/>
          <wp:docPr id="1" name="InsertedImage"/>
          <wp:cNvGraphicFramePr>
            <a:graphicFrameLocks noChangeAspect="1"/>
          </wp:cNvGraphicFramePr>
          <a:graphic>
            <a:graphicData uri="http://schemas.openxmlformats.org/drawingml/2006/picture">
              <pic:pic>
                <pic:nvPicPr>
                  <pic:cNvPr id="0" name="Picture"/>
                  <pic:cNvPicPr/>
                </pic:nvPicPr>
                <pic:blipFill>
                  <a:blip r:embed="{rel_id}"/>
                  <a:stretch>
                    <a:fillRect/>
                  </a:stretch>
                </pic:blipFill>
                <pic:spPr>
                  <a:xfrm>
                    <a:off x="0" y="0"/>
                    <a:ext cx="{cx}" cy="{cy}"/>
                  </a:xfrm>
                  <a:prstGeom prst="rect">
                    <a:avLst/>
                  </a:prstGeom>
                </pic:spPr>
              </pic:pic>
            </a:graphicData>
          </a:graphic>
        </wp:inline>
      </w:drawing>
    </w:r>
    '''
    return " ".join(drawing_xml.split())

def inject_images_into_docx(input_docx, output_docx, placeholder_image_map, text_map=None):
    """
    Injects text replacements (text_map) and image placeholders (placeholder_image_map)
    into the XML parts of input_docx and writes output_docx.
    """
    if not os.path.exists(input_docx):
        raise FileNotFoundError("Input docx not found: " + str(input_docx))

    tmpdir = input_docx + "_tmp_" + uuid.uuid4().hex
    if os.path.exists(tmpdir):
        shutil.rmtree(tmpdir)
    os.makedirs(tmpdir)

    # unzip input_docx
    with zipfile.ZipFile(input_docx, 'r') as zin:
        zin.extractall(tmpdir)

    word_dir = os.path.join(tmpdir, "word")
    ensure_dir(word_dir)
    media_dir = os.path.join(word_dir, "media")
    ensure_dir(media_dir)

    added_media = {}

    # walk XML parts under word/
    for root, dirs, files in os.walk(word_dir):
        for fname in files:
            if not fname.endswith(".xml"):
                continue
            xml_path = os.path.join(root, fname)
            # read raw xml text
            try:
                with open(xml_path, 'rb') as f:
                    txt = f.read().decode('utf-8')
            except Exception:
                continue

            modified = False

            # First do text replacements if requested (useful for textboxes)
            if text_map:
                for tkey, tval in text_map.items():
                    if tkey in txt:
                        txt = txt.replace(tkey, str(tval))
                        modified = True

            for placeholder, img_path in placeholder_image_map.items():
                if not placeholder in txt:
                    continue
                if not img_path or not os.path.exists(img_path):
                    log(f"Image missing for placeholder {placeholder}: {img_path}")
                    continue

                # add image to media folder (once)
                if img_path in added_media:
                    media_fname, w_px, h_px = added_media[img_path]
                else:
                    try:
                        media_fname, w_px, h_px = add_image_file_to_media(tmpdir, img_path)
                    except Exception as e:
                        log(f"Failed adding image to media for {img_path}: {e}")
                        continue
                    added_media[img_path] = (media_fname, w_px, h_px)

                # determine rels file path for this xml part
                xml_rel = os.path.relpath(xml_path, tmpdir)  # e.g., word/document.xml or word/header1.xml
                rels_path = os.path.join(os.path.dirname(xml_rel), "_rels", os.path.basename(xml_rel) + ".rels")
                rels_full = os.path.join(tmpdir, rels_path)
                ensure_dir(os.path.dirname(rels_full))
                if not os.path.exists(rels_full):
                    # create minimal rels file
                    root_rels = etree.Element("{http://schemas.openxmlformats.org/package/2006/relationships}Relationships")
                    tree = etree.ElementTree(root_rels)
                    tree.write(rels_full, xml_declaration=True, encoding="utf-8")

                # add relationship
                try:
                    rId = add_image_relationship(rels_full, media_fname)
                except Exception as e:
                    log(f"Failed to add relationship for media {media_fname}: {e}")
                    continue

                # compute EMU extents
                cx = int(w_px * EMU_PER_PIXEL)
                cy = int(h_px * EMU_PER_PIXEL)

                drawing_snippet = build_drawing_xml(rId, cx, cy)

                # replace placeholder string with drawing XML
                txt = txt.replace(placeholder, drawing_snippet)
                modified = True

            if modified:
                try:
                    with open(xml_path, 'wb') as f:
                        f.write(txt.encode('utf-8'))
                except Exception as e:
                    log(f"Failed to write xml part {xml_path}: {e}")

    # rezip
    with zipfile.ZipFile(output_docx, 'w', zipfile.ZIP_DEFLATED) as zout:
        for foldername, subfolders, filenames in os.walk(tmpdir):
            for filename in filenames:
                filepath = os.path.join(foldername, filename)
                arcname = os.path.relpath(filepath, tmpdir)
                zout.write(filepath, arcname)

    shutil.rmtree(tmpdir)
    return output_docx

# -------------------------
# create partial report: Option B flow (text placeholders first, then XML injection of images)
# -------------------------
def create_partial_report_with_shift_signs(date_str):
    log(f"Creating partial report for {date_str}")

    sign_map = find_shift_sign_photos(date_str)

    # Step 1: load original template (do not try to replace (date) using python-docx)
    try:
        doc = Document(TEMPLATE_ORIG)
    except Exception as e:
        log(f"Failed to open template {TEMPLATE_ORIG}: {e}")
        return None

    # We will perform date replacements at XML level because (date) lives inside a textbox
    human_date = datetime.strptime(date_str, "%Y-%m-%d").strftime("%d %B %Y")

    # --- IMPORTANT FIX: replace (date_with_month) using python-docx BEFORE saving tmp docx ---
    mapping_docx = {
        "(date_with_month)": human_date
    }
    try:
        replace_text_placeholders(doc, mapping_docx)
    except Exception as e:
        log(f"Warning: replace_text_placeholders failed on python-docx layer: {e}")

    # Save temp docx (python-docx changes now saved)
    tmp_text_docx = os.path.join(OUTPUT_DIR, f"temp_text_{date_str}_{uuid.uuid4().hex}.docx")
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    try:
        doc.save(tmp_text_docx)
    except Exception as e:
        log(f"Failed to save temporary text document: {e}")
        return None

    # Step 2: build placeholder->image mapping for injection
    placeholder_image_map = {
        "(shift_1_signin)": sign_map.get("shift_1_signin"),
        "(shift_1_signout)": sign_map.get("shift_1_signout"),
        "(shift_2_signin)": sign_map.get("shift_2_signin"),
        "(shift_2_signout)": sign_map.get("shift_2_signout"),
    }

    # ⭐⭐⭐ IMAGE RESIZING (162 × 162 px) ⭐⭐⭐
    final_image_map = {}
    for placeholder, img_path in placeholder_image_map.items():
        if img_path:
            # save the resized file next to original (with _162 suffix)
            base, ext = os.path.splitext(img_path)
            resized_path = base + "_162.jpg"  # ensure jpg suffix for our resized version
            try:
                resize_image_fixed(img_path, resized_path, 162, 162)
                final_image_map[placeholder] = resized_path
                log(f"Resized image for {placeholder} -> {resized_path}")
            except Exception as e:
                log(f"Resize failed for {img_path}: {e}")
                # fallback to original (best effort)
                final_image_map[placeholder] = img_path
        else:
            final_image_map[placeholder] = None

    placeholder_image_map = final_image_map
    # ⭐⭐⭐ END OF RESIZE LOGIC ⭐⭐⭐

    # Step 3: inject images and also text placeholders into the temp_text_docx, producing final docx
    final_docx = os.path.join(OUTPUT_DIR, f"Daily_Report_{date_str}_partial.docx")
    final_docx_safe = safe_save_docx(final_docx)

    # build XML-level text map (for textboxes and any other XML-only content)
    mapping_for_xml = {
        "(date)": date_str,
        "(date_with_month)": human_date
    }

    try:
        inject_images_into_docx(tmp_text_docx, final_docx_safe, placeholder_image_map, text_map=mapping_for_xml)
    except Exception as e:
        log(f"XML injection failed: {e}")
        # try fallback: just copy tmp_text_docx -> final
        try:
            shutil.copy2(tmp_text_docx, final_docx_safe)
        except Exception as e2:
            log(f"Failed fallback copy: {e2}")
            return None

    # apply final formatting
    try:
        doc_final = Document(final_docx_safe)
        force_arial(doc_final)
        doc_final.save(final_docx_safe)
    except Exception as e:
        log(f"Post python-docx formatting failed: {e}")

    log(f"Saved partial: {final_docx_safe}")

    # cleanup tmp_text_docx
    try:
        if os.path.exists(tmp_text_docx):
            os.remove(tmp_text_docx)
    except:
        pass

    return final_docx_safe

# -------------------------
# REPORT FINALIZATION HELPERS
# -------------------------
def check_report_ready(date_str):
    """
    Return True if:
      - there exists an end_shift record with shift == "2" and its photo file is present
      - AND every record with type == "record_update" has a 'photo' and that photo file exists
    """
    # load local JSON records
    records = load_day_records_local(date_str)
    if not records:
        log(f"No records found locally for {date_str} — cannot finalize.")
        return False

    photos_dir = os.path.join(LOCAL_DIR, date_str, "photos")

    # check for shift 2 end record with photo present
    found_shift2_end = False
    for r in records:
        if r.get("type") == "end_shift" and str(r.get("shift", "")) == "2":
            photo = r.get("photo")
            if photo:
                if os.path.exists(os.path.join(photos_dir, photo)):
                    found_shift2_end = True
                    break
                else:
                    log(f"Shift2 end record references photo {photo} but file missing.")
                    return False
            else:
                log("Shift2 end record has no photo field.")
                return False

    if not found_shift2_end:
        log("Shift 2 end not found yet — not ready.")
        return False

    # ensure all cage update records have photos
    for r in records:
        if r.get("type") == "record_update":
            photo = r.get("photo")
            if not photo:
                log(f"Record {r.get('_filename','?')} missing photo field — not ready.")
                return False
            if not os.path.exists(os.path.join(photos_dir, photo)):
                log(f"Record {r.get('_filename','?')} references missing photo {photo} — not ready.")
                return False

    # all checks passed
    log(f"All checks passed — {date_str} is ready to finalize.")
    return True


def finalize_report(date_str, partial_docx_path=None):
    """
    Move or save the partial_docx_path to final folder and mark in download_db.
    Creates reports/final/ and writes Daily_Report_<date>_FINAL.docx.
    Returns final_path or None on error.
    """
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    final_dir = os.path.join(OUTPUT_DIR, "final")
    os.makedirs(final_dir, exist_ok=True)

    # determine source partial path if not provided (use existing naming convention)
    if partial_docx_path is None:
        partial_docx_path = os.path.join(OUTPUT_DIR, f"Daily_Report_{date_str}_partial.docx")

    if not os.path.exists(partial_docx_path):
        log(f"Partial doc not found to finalize: {partial_docx_path}")
        return None

    base_name = f"Daily_Report_{date_str}_FINAL.docx"
    final_path = os.path.join(final_dir, base_name)

    # if final_path exists, create a version
    if os.path.exists(final_path):
        idx = 2
        while True:
            candidate = os.path.join(final_dir, f"Daily_Report_{date_str}_FINAL_v{idx}.docx")
            if not os.path.exists(candidate):
                final_path = candidate
                break
            idx += 1

    try:
        # attempt to move (rename). If locked, copy + keep partial.
        try:
            os.replace(partial_docx_path, final_path)  # atomic where possible
            log(f"Moved partial to final: {final_path}")
        except PermissionError:
            # file locked - copy instead
            shutil.copy2(partial_docx_path, final_path)
            log(f"Partial file locked; copied to final path: {final_path}")
    except Exception as e:
        log(f"Failed to finalize report: {e}")
        return None

    # mark finalized in download_db so future syncs skip this day
    if date_str not in download_db:
        download_db[date_str] = {}

    download_db[date_str]["finalized"] = True
    download_db[date_str]["finalized_at"] = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    save_download_db(download_db)

    return final_path


# -------------------------
# Sync one day
# -------------------------
def sync_day(day):
    global download_db
    new_data = False
    new_photos = False


    res = safe_request(BASE_URL + f"{day}/list")
    if res is None:
        log(f"Could not fetch file list for {day}")
        return

    files = res.json()
    server_data = files.get("data", [])
    server_photos = files.get("photos", [])

    if day not in download_db:
        download_db[day] = {"data": [], "photos": []}

    known_data = set(download_db[day]["data"])
    known_photos = set(download_db[day]["photos"])

    data_dir = os.path.join(LOCAL_DIR, day, "data")
    photos_dir = os.path.join(LOCAL_DIR, day, "photos")
    os.makedirs(data_dir, exist_ok=True)
    os.makedirs(photos_dir, exist_ok=True)

    # fetch JSONs
    for f in server_data:
        if f not in known_data:
            if download_file(BASE_URL + f"{day}/data/{f}", os.path.join(data_dir, f)):
                download_db[day]["data"].append(f)
                save_download_db(download_db)
                new_data = True

    # fetch photos
    for f in server_photos:
        if f not in known_photos:
            if download_file(BASE_URL + f"{day}/photos/{f}", os.path.join(photos_dir, f)):
                download_db[day]["photos"].append(f)
                save_download_db(download_db)
                new_photos = True

    # create report after sync
    if new_data or new_photos:
        partial_path = create_partial_report_with_shift_signs(day)
        if partial_path is None:
            log("Partial report creation failed.")
        else:
            # If this day is already finalized in DB, skip any further processing
            if download_db.get(day, {}).get("finalized"):
                log(f"{day} already finalized — skipping finalization.")
            else:
                # Check finalization criteria
                if check_report_ready(day):
                    final_path = finalize_report(day, partial_docx_path=partial_path)
                    if final_path:
                        log(f"Report finalized: {final_path}")
                    else:
                        log("Finalization attempt failed.")
    else:
        log("No new files – report unchanged.")


    # delete on server if threshold reached
    if len(server_data) + len(server_photos) >= 10:
        log(f"Reached limit, deleting server files for {day}")
        delete_from_server(day)

# -------------------------
# Get available date folders on server
# -------------------------
def get_available_dates():
    res = safe_request(BASE_URL + "list_dates")
    if res is None:
        log("Could not get date folder list")
        return []
    try:
        return res.json()
    except Exception:
        return []

# -------------------------
# Main loop
# -------------------------
def main_loop():
    log("=== SYNC SYSTEM STARTED ===")
    while True:
        log("Checking for new files...")
        dates = get_available_dates()
        if dates:
            for day in dates:
                sync_day(day)
        log(f"Cycle complete. Sleeping {LOOP_INTERVAL} sec...")
        time.sleep(LOOP_INTERVAL)

if __name__ == "__main__":
    main_loop()
